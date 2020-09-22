#Crate a separate word file for each row in excel and store particular cell values as a table in word file
import openpyxl as xl
from docx import Document

path2 = 'book1.xlsx'

src_wb = xl.load_workbook('book1.xlsx')
src_ws = src_wb.worksheets[4]
i=0
for row in src_ws:
	fn="Student " + str(i)+".docx"

	document = Document()
	document.add_heading('Student Details', 0)
	table = document.add_table(rows=8, cols=2)
	cell = table.cell(0, 0)
	cell.text = 'Name'
	cell = table.cell(1, 0)
	cell.text = 'Registration Number'
	cell = table.cell(2, 0)
	cell.text = 'Email ID'
	cell = table.cell(3, 0)
	cell.text = 'Contact'
	cell = table.cell(4, 0)
	cell.text = 'BTech'
	cell = table.cell(5, 0)
	cell.text = '12th'
	cell = table.cell(6, 0)
	cell.text = '10th'
	cell = table.cell(7, 0)
	cell.text = 'Placement/Higher Studies'


	cell = table.cell(0, 1)
	cell.text = str(row[2].value)
	cell = table.cell(1, 1)
	cell.text = str(row[1].value)
	cell = table.cell(2, 1)
	cell.text = str(row[5].value)
	cell = table.cell(3, 1)
	cell.text = str(row[6].value)
	cell = table.cell(4, 1)
	cell.text = str(row[9].value)
	cell = table.cell(5, 1)
	cell.text = str(row[8].value)
	cell = table.cell(6, 1)
	cell.text = str(row[7].value)
	cell = table.cell(7, 1)
	cell.text = 'Not yet placed'

	document.save(fn)
	i+=1
	if(i==70):
		break